"""DOCX report generator via python-docx.

Produces genuinely editable Word documents with professional formatting.
All content is derived from persisted analysis artifacts.
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from app.schemas.publication_export import PublicationExportConfig, FigureMetadata


def _add_cover_page(doc: Document, config: PublicationExportConfig) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        p.add_run("\n")

    title_run = p.add_run(config.title)
    title_run.font.size = Pt(24)
    title_run.bold = True

    if config.subtitle:
        p.add_run("\n\n")
        sub = p.add_run(config.subtitle)
        sub.font.size = Pt(14)
        sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if config.author_name:
        p.add_run("\n\n\n")
        author = p.add_run(config.author_name)
        author.font.size = Pt(12)

    if config.institution_name:
        p.add_run("\n")
        inst = p.add_run(config.institution_name)
        inst.font.size = Pt(11)
        inst.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if config.course_or_supervisor:
        p.add_run("\n")
        sup = p.add_run(config.course_or_supervisor)
        sup.font.size = Pt(11)
        sup.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if config.report_date:
        p.add_run("\n\n")
        dt = p.add_run(config.report_date.strftime("%B %d, %Y"))
        dt.font.size = Pt(11)

    doc.add_page_break()


def _parse_md_table(table_text: str) -> tuple[list[str], list[list[str]]]:
    lines = [l.strip() for l in table_text.strip().split("\n") if l.strip().startswith("|")]
    if len(lines) < 2:
        return [], []
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[1:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.fullmatch(r"-+", c.replace(":", "")) for c in cells if c):
            continue
        rows.append(cells)
    return headers, rows


def _add_md_table_to_doc(doc: Document, table_text: str) -> None:
    headers, rows = _parse_md_table(table_text)
    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < len(headers):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)


def _add_markdown_content(doc: Document, md: str, figures_dir: Path | None = None) -> None:
    in_table = False
    table_lines: list[str] = []

    for line in md.split("\n"):
        stripped = line.strip()

        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            _add_md_table_to_doc(doc, "\n".join(table_lines))
            table_lines = []
            in_table = False

        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("**Table:"):
            title = stripped.strip("*").strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(10)
        elif stripped.startswith("*Note:") or stripped.startswith("*Consistency"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*").strip())
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
        else:
            bold_re = re.compile(r"\*\*(.+?)\*\*")
            parts = bold_re.split(stripped)
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    if in_table and table_lines:
        _add_md_table_to_doc(doc, "\n".join(table_lines))


def generate_docx(
    markdown_content: str,
    config: PublicationExportConfig,
    figures: list[FigureMetadata] | None = None,
    figures_dir: Path | None = None,
    output_path: Path | None = None,
) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    if config.include_cover_page:
        _add_cover_page(doc, config)

    _add_markdown_content(doc, markdown_content, figures_dir)

    if figures and figures_dir and config.include_figures:
        doc.add_page_break()
        doc.add_heading("Figures", level=2)
        for i, fig in enumerate(figures, 1):
            fig_path = figures_dir / fig.filename
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(5.5))
                cap = doc.add_paragraph()
                run = cap.add_run(f"Figure {i}: {fig.caption}")
                run.italic = True
                run.font.size = Pt(9)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = disclaimer.add_run(
        "Disclaimer: The reported estimates describe statistical associations "
        "and should not be interpreted as causal effects without additional "
        "identification assumptions."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if output_path:
        doc.save(str(output_path))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
